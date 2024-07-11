from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from .models import Post, Form
from .forms import PostForm, SaveForm
from docx import Document
from django_sendfile import sendfile
from docx.shared import Inches

def post_list(request):
    posts = Post.objects.filter(published_date__lte=timezone.now()).order_by('published_date')
    return render(request, 'blog/post_list.html', {'posts': posts})

def post_detail(request, pk):
    post = get_object_or_404(Post, pk=pk)
    return render(request, 'blog/post_detail.html', {'post': post})

def post_new(request):
    if request.method == "POST":
        form = PostForm(request.POST)
        if form.is_valid():
            post = form.save(commit=False)
            try:
                post.author = request.user
            except:
                return redirect('post_detail', pk=post.pk)  
            post.published_date = timezone.now()
            post.save()
            return redirect('post_detail', pk=post.pk)
    else:
        form = PostForm()
    return render(request, 'blog/post_edit.html', {'form': form})

def post_edit(request, pk):
    post = get_object_or_404(Post, pk=pk)
    if request.method == "POST":
        form = PostForm(request.POST, instance=post)
        if form.is_valid():
            post = form.save(commit=False)
            post.author = request.user
            post.published_date = timezone.now()
            post.save()
            return redirect('post_detail', pk=post.pk)
    else:
        form = PostForm(instance=post)
    return render(request, 'blog/post_edit.html', {'form': form})

def create_form(request):
    context = {}
    context['form'] = SaveForm()
    return render(request, "blog/show_form.html", context)

def form_list(request):
    posts = Form.objects.all()
    return render(request, 'blog/form_list.html', {'forms': posts})

def form_detail(request, pk):
    form = get_object_or_404(Form, pk=pk)
    return render(request, 'blog/form_detail.html', {'form': form})

def form_edit(request, pk):
    post = get_object_or_404(Form, pk=pk)
    if request.method == "POST":
        form = SaveForm(request.POST, instance=post)
        if form.is_valid():
            post = form.save(commit=False)
            post.save()
            return redirect('form_detail', pk=post.pk)
    else:
        form = SaveForm(instance=post)
    return render(request, 'blog/form_edit.html', {'form': form})




def run(name, address_L1, address_L2, postcode):
    document = Document()
    document.add_heading(f'Form for {name}', 0)
    document.add_paragraph(f'NAME: {name}')
    document.add_paragraph(f'ADDRESS LINE 1: {address_L1}')
    document.add_paragraph(f'ADDRESS LINE 2: {address_L2}')
    document.add_paragraph(f'POSTCODE: {postcode}')
    document.add_page_break()
    foldername = "media-private"
    filename = "form.docx"
    document.save(f"{foldername}/{filename}")
    return filename

def download_form(request, pk):
    post = get_object_or_404(Form, pk=pk)
    filename = run(post.name, post.address_L1, post.address_L2, post.postcode)
    return sendfile(request, filename, attachment=True)
