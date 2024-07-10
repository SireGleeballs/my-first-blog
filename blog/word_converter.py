
from docx import Document
from docx.shared import Inches

def run(name, address_L1, address_L2, postcode):
    document = Document()

    document.add_heading(f'Form for {name}', 0)

    document.add_paragraph(f'NAME: {name}')
    document.add_paragraph(f'ADDRESS LINE 1: {address_L1}')
    document.add_paragraph(f'ADDRESS LINE 2: {address_L2}')
    document.add_paragraph(f'POSTCODE: {postcode}')


    document.add_page_break()
    document.save('demo.docx')

run("Daniel Spence", "L1XXXXX", "L2XXXXX", "PCXXXXX")

